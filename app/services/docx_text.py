"""Extract plain text from .docx (Word) bytes for document import."""

from io import BytesIO


def extract_plain_text_from_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(BytesIO(data))
    lines: list[str] = []
    for para in doc.paragraphs:
        lines.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            lines.append(
                "\t".join((cell.text or "").replace("\n", " ") for cell in row.cells)
            )
    return "\n".join(lines).strip()
